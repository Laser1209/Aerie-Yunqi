"""Aerie v13.9.8 · 文档写作模块

支持 5 类文档模板：
  - diary: 日记
  - report: 报告
  - spec: 技术规格
  - research: 研究报告
  - resume: 简历

支持 4 种导出格式：
  - Markdown (.md)
  - HTML (.html)
  - PDF (需要 WeasyPrint 或浏览器打印，回退 HTML)
  - Word (.docx, 需要 python-docx)
"""

from __future__ import annotations
import time
import json
import logging
from pathlib import Path
from typing import Any, Optional
from dataclasses import dataclass, field
from enum import Enum

logger = logging.getLogger(__name__)


class DocType(str, Enum):
    """文档类型"""
    DIARY = "diary"
    REPORT = "report"
    SPEC = "spec"
    RESEARCH = "research"
    RESUME = "resume"


class ExportFormat(str, Enum):
    """导出格式"""
    MARKDOWN = "md"
    HTML = "html"
    PDF = "pdf"
    DOCX = "docx"


DOC_TYPE_NAMES = {
    DocType.DIARY: "日记",
    DocType.REPORT: "报告",
    DocType.SPEC: "技术规格",
    DocType.RESEARCH: "研究报告",
    DocType.RESUME: "简历",
}


# 文档模板
DOC_TEMPLATES = {
    DocType.DIARY: """# {title}

> **日期**: {date}
> **天气**: {weather}
> **心情**: {mood}

---

## 今日点滴

{content}

---

## 感悟与收获

{reflection}

---

## 明日计划

{tomorrow}

---
*由 Aerie 记录 · {timestamp}*
""",

    DocType.REPORT: """# {title}

| 项目 | 内容 |
|---|---|
| 报告类型 | {report_type} |
| 编写日期 | {date} |
| 编写人 | {author} |
| 版本 | v{version} |

---

## 摘要

{summary}

---

## 一、背景

{background}

---

## 二、详细内容

{content}

---

## 三、数据分析

{data_analysis}

---

## 四、结论与建议

{conclusion}

---

## 五、附录

{appendix}

---
*{author} · {date}*
""",

    DocType.SPEC: """# {title} - 技术规格文档

| 字段 | 值 |
|---|---|
| 文档版本 | v{version} |
| 创建日期 | {date} |
| 作者 | {author} |
| 状态 | {status} |

---

## 1. 概述

### 1.1 目的

{purpose}

### 1.2 范围

{scope}

### 1.3 术语定义

{glossary}

---

## 2. 系统架构

{architecture}

---

## 3. 功能规格

### 3.1 功能清单

{features}

### 3.2 详细规格

{detailed_spec}

---

## 4. 接口定义

{interfaces}

---

## 5. 非功能需求

### 5.1 性能要求

{performance}

### 5.2 安全要求

{security}

### 5.3 可靠性

{reliability}

---

## 6. 实施计划

{implementation_plan}

---

## 7. 附录

{appendix}

---
*技术规格文档 · {author} · v{version}*
""",

    DocType.RESEARCH: """# {title}

## 研究报告

| 项目 | 内容 |
|---|---|
| 研究主题 | {topic} |
| 研究员 | {author} |
| 开始日期 | {start_date} |
| 完成日期 | {end_date} |
| 版本 | v{version} |

---

## 摘要

{abstract}

**关键词**: {keywords}

---

## 1. 研究背景

{background}

---

## 2. 研究方法

{methodology}

---

## 3. 研究内容

{content}

---

## 4. 实验与数据分析

{experiments}

---

## 5. 研究结论

{conclusion}

---

## 6. 展望与后续工作

{future_work}

---

## 参考文献

{references}

---
*研究报告 · {author} · {end_date}*
""",

    DocType.RESUME: """# {name}

## 个人简介

{summary}

---

### 基本信息

| 项目 | 内容 |
|---|---|
| 姓名 | {name} |
| 职位 | {position} |
| 邮箱 | {email} |
| 电话 | {phone} |
| 所在地 | {location} |

---

## 工作经历

{experience}

---

## 教育背景

{education}

---

## 技能专长

{skills}

---

## 项目经历

{projects}

---

## 其他

{other}

---
*{name} · 个人简历*
""",
}


# 默认填充值
DEFAULT_FIELDS = {
    DocType.DIARY: {
        "title": "今日日记",
        "date": "",
        "weather": "晴",
        "mood": "愉快",
        "content": "今天是美好的一天。",
        "reflection": "学到了很多新东西。",
        "tomorrow": "继续努力！",
    },
    DocType.REPORT: {
        "title": "工作报告",
        "report_type": "周报",
        "date": "",
        "author": "Aerie",
        "version": "1.0",
        "summary": "本报告概述了近期工作进展。",
        "background": "根据项目计划安排...",
        "content": "详细内容...",
        "data_analysis": "数据分析...",
        "conclusion": "结论与建议...",
        "appendix": "附录内容...",
    },
    DocType.SPEC: {
        "title": "系统名称",
        "version": "1.0",
        "date": "",
        "author": "Aerie",
        "status": "初稿",
        "purpose": "本文档描述...",
        "scope": "本文档涵盖...",
        "glossary": "- **术语1**: 定义...",
        "architecture": "系统架构图及说明...",
        "features": "- 功能1\n- 功能2",
        "detailed_spec": "详细规格描述...",
        "interfaces": "API 接口定义...",
        "performance": "性能指标...",
        "security": "安全要求...",
        "reliability": "可靠性要求...",
        "implementation_plan": "实施计划...",
        "appendix": "附录...",
    },
    DocType.RESEARCH: {
        "title": "研究报告",
        "topic": "研究主题",
        "author": "研究员",
        "start_date": "",
        "end_date": "",
        "version": "1.0",
        "abstract": "研究摘要...",
        "keywords": "关键词1, 关键词2",
        "background": "研究背景...",
        "methodology": "研究方法...",
        "content": "研究内容...",
        "experiments": "实验设计与结果...",
        "conclusion": "研究结论...",
        "future_work": "后续工作方向...",
        "references": "[1] 参考文献1\n[2] 参考文献2",
    },
    DocType.RESUME: {
        "name": "姓名",
        "summary": "个人简介...",
        "position": "目标职位",
        "email": "email@example.com",
        "phone": "138-xxxx-xxxx",
        "location": "城市",
        "experience": "### 公司A · 职位 · 2020-至今\n\n- 工作内容1\n- 工作内容2",
        "education": "### 学校名称 · 专业 · 学位 · 2016-2020\n\n- GPA: 3.8/4.0\n- 荣誉：一等奖学金",
        "skills": "- 编程语言: Python, JavaScript, Go\n- 框架: Vue, React, Django\n- 其他: Docker, Kubernetes",
        "projects": "### 项目名称\n\n- 项目描述...\n- 技术栈...\n- 我的贡献...",
        "other": "兴趣爱好、证书等...",
    },
}


@dataclass
class Document:
    """文档对象"""
    doc_type: DocType
    title: str
    fields: dict = field(default_factory=dict)
    content: str = ""
    created_at: float = field(default_factory=time.time)
    updated_at: float = field(default_factory=time.time)

    @property
    def type_name(self) -> str:
        return DOC_TYPE_NAMES.get(self.doc_type, self.doc_type.value)

    def render_markdown(self) -> str:
        """渲染为 Markdown"""
        template = DOC_TEMPLATES.get(self.doc_type, "# {title}\n\n{content}")
        # 合并默认值和自定义值
        data = DEFAULT_FIELDS.get(self.doc_type, {}).copy()
        data.update(self.fields)
        data["title"] = self.title
        data["content"] = self.content or data.get("content", "")

        # 日期填充
        if not data.get("date"):
            data["date"] = time.strftime("%Y-%m-%d")
        if not data.get("start_date"):
            data["start_date"] = time.strftime("%Y-%m-%d")
        if not data.get("end_date"):
            data["end_date"] = time.strftime("%Y-%m-%d")
        data["timestamp"] = time.strftime("%Y-%m-%d %H:%M:%S")

        try:
            return template.format(**data)
        except KeyError as e:
            # 缺少字段时回退到简单格式
            logger.warning(f"模板缺少字段 {e}，回退到简单格式")
            return f"# {self.title}\n\n{self.content}"

    def render_html(self, style: str = "default") -> str:
        """渲染为 HTML

        Args:
            style: 样式主题 - default/elegant/minimal
        """
        md_content = self.render_markdown()
        html_body = self._markdown_to_html(md_content)

        css = self._get_css(style)

        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.title}</title>
    <style>
{css}
    </style>
</head>
<body>
    <div class="document">
        {html_body}
    </div>
</body>
</html>"""

    def _markdown_to_html(self, md: str) -> str:
        """简单 Markdown 转 HTML（无需第三方依赖）

        支持：标题、段落、列表、表格、分隔线、引用、加粗、代码块
        """
        lines = md.split("\n")
        html_lines = []
        in_code_block = False
        in_table = False
        in_list = False
        list_type = "ul"

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # 代码块
            if stripped.startswith("```"):
                if in_code_block:
                    html_lines.append("</code></pre>")
                    in_code_block = False
                else:
                    html_lines.append("<pre><code>")
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                html_lines.append(self._escape_html(line))
                i += 1
                continue

            # 分隔线
            if stripped in ("---", "***", "___"):
                html_lines.append("<hr>")
                i += 1
                continue

            # 标题
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                text = stripped[level:].strip()
                html_lines.append(f"<h{level}>{self._inline_format(text)}</h{level}>")
                i += 1
                continue

            # 引用
            if stripped.startswith("> "):
                quote_lines = []
                while i < len(lines) and lines[i].strip().startswith("> "):
                    quote_lines.append(lines[i].strip()[2:])
                    i += 1
                html_lines.append(f"<blockquote>{self._inline_format(' '.join(quote_lines))}</blockquote>")
                continue

            # 表格
            if "|" in stripped and i + 1 < len(lines) and "---" in lines[i + 1]:
                # 表头
                header_cells = [c.strip() for c in stripped.split("|") if c.strip()]
                i += 2  # 跳过分隔行

                rows_html = [f"<th>{self._inline_format(c)}</th>" for c in header_cells]
                html_lines.append("<table><thead><tr>" + "".join(rows_html) + "</tr></thead><tbody>")

                # 表体
                while i < len(lines) and "|" in lines[i].strip() and lines[i].strip():
                    cells = [c.strip() for c in lines[i].strip().split("|") if c.strip()]
                    row_html = [f"<td>{self._inline_format(c)}</td>" for c in cells]
                    html_lines.append("<tr>" + "".join(row_html) + "</tr>")
                    i += 1

                html_lines.append("</tbody></table>")
                continue

            # 无序列表
            if stripped.startswith(("- ", "* ", "+ ")):
                if not in_list or list_type != "ul":
                    if in_list:
                        html_lines.append(f"</{list_type}>")
                    html_lines.append("<ul>")
                    in_list = True
                    list_type = "ul"
                item_text = stripped[2:].strip()
                html_lines.append(f"<li>{self._inline_format(item_text)}</li>")
                i += 1
                continue

            # 有序列表
            if stripped and stripped[0].isdigit() and ". " in stripped[:5]:
                if not in_list or list_type != "ol":
                    if in_list:
                        html_lines.append(f"</{list_type}>")
                    html_lines.append("<ol>")
                    in_list = True
                    list_type = "ol"
                item_text = stripped[stripped.index(". ") + 2:].strip()
                html_lines.append(f"<li>{self._inline_format(item_text)}</li>")
                i += 1
                continue

            # 空行 - 结束列表
            if not stripped:
                if in_list:
                    html_lines.append(f"</{list_type}>")
                    in_list = False
                html_lines.append("")
                i += 1
                continue

            # 普通段落
            if in_list:
                html_lines.append(f"</{list_type}>")
                in_list = False
            html_lines.append(f"<p>{self._inline_format(stripped)}</p>")
            i += 1

        if in_list:
            html_lines.append(f"</{list_type}>")
        if in_code_block:
            html_lines.append("</code></pre>")

        return "\n".join(html_lines)

    def _inline_format(self, text: str) -> str:
        """行内格式化：加粗、斜体、代码、链接"""
        import re
        # 转义 HTML
        text = self._escape_html(text)
        # 加粗
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        # 斜体
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        # 行内代码
        text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
        return text

    def _escape_html(self, text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    def _get_css(self, style: str = "default") -> str:
        """获取 CSS 样式"""
        if style == "minimal":
            return """
body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; line-height: 1.6; color: #333; max-width: 800px; margin: 0 auto; padding: 40px 20px; }
h1, h2, h3 { font-weight: 600; }
h1 { border-bottom: 2px solid #eee; padding-bottom: 10px; }
hr { border: none; border-top: 1px solid #eee; margin: 30px 0; }
blockquote { border-left: 4px solid #ddd; margin: 0; padding-left: 16px; color: #666; }
table { border-collapse: collapse; width: 100%; margin: 20px 0; }
th, td { border: 1px solid #ddd; padding: 12px; text-align: left; }
th { background: #f8f9fa; }
code { background: #f5f5f5; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; }
pre { background: #f5f5f5; padding: 16px; border-radius: 6px; overflow-x: auto; }
"""
        elif style == "elegant":
            return """
@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@400;600&display=swap');
body { font-family: 'Noto Serif SC', Georgia, serif; line-height: 1.8; color: #2c3e50; max-width: 850px; margin: 0 auto; padding: 50px 30px; background: #fafbfc; }
.document { background: white; padding: 60px; box-shadow: 0 2px 20px rgba(0,0,0,0.06); border-radius: 4px; }
h1 { color: #1a365d; font-size: 2em; border-bottom: 3px solid #667eea; padding-bottom: 12px; }
h2 { color: #2d3748; font-size: 1.5em; margin-top: 30px; }
h3 { color: #4a5568; font-size: 1.2em; }
hr { border: none; border-top: 2px solid #edf2f7; margin: 30px 0; }
blockquote { border-left: 4px solid #667eea; background: #f7fafc; margin: 20px 0; padding: 16px 20px; color: #4a5568; border-radius: 0 4px 4px 0; }
table { border-collapse: collapse; width: 100%; margin: 20px 0; font-size: 0.95em; }
th, td { border: 1px solid #e2e8f0; padding: 14px 16px; text-align: left; }
th { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; font-weight: 500; }
tr:nth-child(even) { background: #f7fafc; }
code { background: #edf2f7; padding: 2px 8px; border-radius: 4px; font-family: 'Consolas', monospace; font-size: 0.9em; color: #e53e3e; }
pre { background: #1a202c; color: #e2e8f0; padding: 20px; border-radius: 8px; overflow-x: auto; }
pre code { background: none; color: inherit; padding: 0; }
strong { color: #2d3748; }
"""
        else:  # default
            return """
body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC', 'Microsoft YaHei', sans-serif; line-height: 1.7; color: #1f2937; max-width: 800px; margin: 0 auto; padding: 40px 24px; background: #f9fafb; }
.document { background: white; padding: 48px; border-radius: 8px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }
h1 { font-size: 1.8em; color: #111827; border-bottom: 2px solid #3b82f6; padding-bottom: 12px; margin-top: 0; }
h2 { font-size: 1.4em; color: #1f2937; margin-top: 32px; padding-bottom: 8px; border-bottom: 1px solid #e5e7eb; }
h3 { font-size: 1.15em; color: #374151; margin-top: 24px; }
hr { border: none; border-top: 1px solid #e5e7eb; margin: 28px 0; }
blockquote { border-left: 4px solid #3b82f6; background: #eff6ff; margin: 16px 0; padding: 12px 20px; color: #1e40af; border-radius: 0 6px 6px 0; }
table { border-collapse: collapse; width: 100%; margin: 20px 0; font-size: 0.95em; }
th, td { border: 1px solid #e5e7eb; padding: 12px 16px; text-align: left; }
th { background: #f3f4f6; font-weight: 600; color: #1f2937; }
tr:nth-child(even) { background: #f9fafb; }
code { background: #f3f4f6; padding: 2px 6px; border-radius: 4px; font-family: 'Consolas', 'Monaco', monospace; font-size: 0.88em; color: #dc2626; }
pre { background: #1f2937; color: #f9fafb; padding: 20px; border-radius: 6px; overflow-x: auto; }
pre code { background: none; color: inherit; padding: 0; }
ul, ol { padding-left: 24px; }
li { margin: 6px 0; }
strong { color: #111827; }
"""

    def to_dict(self) -> dict:
        return {
            "doc_type": self.doc_type.value,
            "title": self.title,
            "fields": self.fields,
            "content": self.content,
            "created_at": self.created_at,
            "updated_at": self.updated_at,
        }


class DocWriter:
    """文档写作器

    支持 5 类文档模板，4 种导出格式。
    """

    def __init__(self, output_dir: str = "data/documents"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._check_dependencies()

    def _check_dependencies(self) -> dict:
        """检查可选依赖"""
        deps = {
            "weasyprint": False,
            "python_docx": False,
            "markdown": False,
        }
        try:
            import weasyprint  # noqa: F401
            deps["weasyprint"] = True
        except ImportError:
            pass
        try:
            import docx  # noqa: F401
            deps["python_docx"] = True
        except ImportError:
            pass
        try:
            import markdown  # noqa: F401
            deps["markdown"] = True
        except ImportError:
            pass
        return deps

    def create_document(self, doc_type: DocType, title: str,
                        fields: Optional[dict] = None,
                        content: str = "") -> Document:
        """创建文档"""
        return Document(
            doc_type=doc_type,
            title=title,
            fields=fields or {},
            content=content,
        )

    def get_template_fields(self, doc_type: DocType) -> dict:
        """获取模板字段"""
        return DEFAULT_FIELDS.get(doc_type, {}).copy()

    def render(self, doc: Document, fmt: ExportFormat,
               style: str = "default") -> str:
        """渲染文档为指定格式

        Returns:
            渲染后的内容字符串（PDF 返回文件路径）
        """
        if fmt == ExportFormat.MARKDOWN:
            return doc.render_markdown()
        elif fmt == ExportFormat.HTML:
            return doc.render_html(style=style)
        elif fmt == ExportFormat.PDF:
            return self._render_pdf(doc, style)
        elif fmt == ExportFormat.DOCX:
            return self._render_docx(doc)
        else:
            raise ValueError(f"不支持的导出格式: {fmt}")

    def export(self, doc: Document, fmt: ExportFormat,
               filename: Optional[str] = None,
               style: str = "default") -> Path:
        """导出为文件

        Returns:
            导出文件路径
        """
        if not filename:
            safe_title = "".join(c for c in doc.title if c.isalnum() or c in (" ", "-", "_"))
            safe_title = safe_title.strip().replace(" ", "_")
            filename = f"{safe_title}.{fmt.value}"

        filepath = self.output_dir / filename

        if fmt == ExportFormat.PDF:
            # PDF 返回的是文件路径
            content = self._render_pdf(doc, style, filepath)
            return Path(content)
        elif fmt == ExportFormat.DOCX:
            content = self._render_docx(doc, filepath)
            return Path(content)
        else:
            content = self.render(doc, fmt, style)
            filepath.write_text(content, encoding="utf-8")
            return filepath

    def _render_pdf(self, doc: Document, style: str = "default",
                    output_path: Optional[Path] = None) -> str:
        """渲染 PDF

        优先使用 WeasyPrint，不可用时返回 HTML 路径作为回退
        """
        deps = self._check_dependencies()

        if deps["weasyprint"]:
            try:
                from weasyprint import HTML
                html_content = doc.render_html(style=style)
                if output_path is None:
                    output_path = self.output_dir / f"{doc.title}.pdf"

                HTML(string=html_content).write_pdf(str(output_path))
                return str(output_path)
            except Exception as e:
                logger.warning(f"WeasyPrint 渲染 PDF 失败: {e}，回退到 HTML")

        # 回退：保存为 HTML 并提示
        html_path = self.output_dir / f"{doc.title}.html"
        html_content = doc.render_html(style=style)
        html_path.write_text(html_content, encoding="utf-8")
        logger.info("PDF 依赖不可用，已导出为 HTML（可用浏览器打印为 PDF）")
        return str(html_path)

    def _render_docx(self, doc: Document,
                     output_path: Optional[Path] = None) -> str:
        """渲染 Word 文档

        需要 python-docx，不可用时回退到 Markdown
        """
        deps = self._check_dependencies()

        if deps["python_docx"]:
            try:
                from docx import Document as DocxDocument
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                if output_path is None:
                    output_path = self.output_dir / f"{doc.title}.docx"

                docx_doc = DocxDocument()

                # 简单实现：把 Markdown 内容按行解析写入
                md_content = doc.render_markdown()
                lines = md_content.split("\n")

                for line in lines:
                    stripped = line.strip()
                    if not stripped:
                        docx_doc.add_paragraph()
                        continue

                    if stripped.startswith("# "):
                        p = docx_doc.add_heading(stripped[2:], level=1)
                    elif stripped.startswith("## "):
                        docx_doc.add_heading(stripped[3:], level=2)
                    elif stripped.startswith("### "):
                        docx_doc.add_heading(stripped[4:], level=3)
                    elif stripped in ("---", "***", "___"):
                        docx_doc.add_paragraph("─" * 50)
                    elif stripped.startswith("> "):
                        p = docx_doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.5)
                        run = p.add_run(stripped[2:])
                        run.italic = True
                    else:
                        docx_doc.add_paragraph(stripped)

                docx_doc.save(str(output_path))
                return str(output_path)
            except Exception as e:
                logger.warning(f"python-docx 渲染失败: {e}，回退到 Markdown")

        # 回退：保存为 .md 文件
        md_path = self.output_dir / f"{doc.title}.md"
        md_content = doc.render_markdown()
        md_path.write_text(md_content, encoding="utf-8")
        logger.info("DOCX 依赖不可用，已导出为 Markdown")
        return str(md_path)

    def quick_create(self, doc_type: DocType, title: str,
                     content: str = "", **fields) -> Document:
        """快速创建文档"""
        return self.create_document(doc_type, title, fields=fields, content=content)

    def list_documents(self) -> list[Path]:
        """列出已导出的文档"""
        if not self.output_dir.exists():
            return []
        return sorted(self.output_dir.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True)
