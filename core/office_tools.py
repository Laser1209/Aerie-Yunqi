"""Aerie v13.9 — Office Tools 办公场景工具集（25+ 工具）

注册到 tool_registry 供 AI function calling 调用的办公类工具，分五大类：

【文件管理类（8个）】
- document_create     创建文档（Markdown / TXT）
- document_read       读取本地文档
- file_search         本地文件搜索
- directory_list      目录遍历
- file_copy           文件复制
- file_move           文件移动
- file_rename         文件重命名
- directory_create    创建目录

【文档处理类（5个）】
- text_summary        文本摘要
- document_convert    文档格式转换
- word_generate       Word 文档生成（python-docx）
- spreadsheet_analyze 表格数据分析
- csv_generate        CSV 表格生成

【系统操作类（5个）】
- calendar_list       日历事件列表
- calendar_create     创建日历事件
- system_info         系统信息查询
- process_list        进程列表
- app_open            打开应用

【数据分析类（4个）】
- data_stats          数据统计
- data_filter         数据过滤
- data_sort           数据排序
- chart_generate      图表生成（SVG）

【网络工具类（4个）】
- web_fetch           网页抓取
- weather_query       天气查询
- translation         文本翻译
- code_search         代码搜索
"""

from __future__ import annotations

import csv
import os
import sys
import json
import shutil
import platform
import logging
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)

# 办公文件默认保存目录（可通过 settings.office.dir 自定义）
_default_office_dir = Path(os.path.expanduser("~/AerieOffice"))
OFFICE_DIR: Path = _default_office_dir


def get_office_dir() -> Path:
    """获取当前办公文件保存目录。

    优先从 settings.yaml 的 office.dir 读取，未设置则使用默认 ~/AerieOffice。
    每次调用都会重新读取 settings，确保设置变更后立即生效。
    """
    global OFFICE_DIR
    try:
        from config.persona_loader import load_settings
        settings = load_settings()
        custom_dir = (
            settings.get("office", {}).get("dir")
            if isinstance(settings, dict)
            else None
        )
        if custom_dir and str(custom_dir).strip():
            p = Path(os.path.expanduser(str(custom_dir).strip()))
            p.mkdir(parents=True, exist_ok=True)
            OFFICE_DIR = p
            return OFFICE_DIR
    except Exception:
        logger.exception("get_office_dir: 读取 settings 失败，使用默认目录")
    # 兜底：默认目录
    _default_office_dir.mkdir(parents=True, exist_ok=True)
    OFFICE_DIR = _default_office_dir
    return OFFICE_DIR


def set_office_dir(path: str | Path) -> dict:
    """设置办公文件保存目录并持久化到 settings.yaml。

    Args:
        path: 目录路径（支持 ~ 展开）

    Returns:
        {"success": bool, "path": str, "error": str}
    """
    try:
        from config.persona_loader import save_settings
        p = Path(os.path.expanduser(str(path).strip()))
        p.mkdir(parents=True, exist_ok=True)
        save_settings({"office": {"dir": str(p.resolve())}})
        global OFFICE_DIR
        OFFICE_DIR = p.resolve()
        logger.info("office dir set to %s", OFFICE_DIR)
        return {"success": True, "path": str(OFFICE_DIR)}
    except Exception as e:
        logger.exception("set_office_dir error")
        return {"success": False, "error": str(e)}


# 启动时初始化一次
get_office_dir()


# ── 工具函数 ──────────────────────────────────────


def tool_document_create(
    filename: str,
    content: str,
    format: str = "markdown",
) -> dict:
    """创建一个办公文档并保存到本地。

    Args:
        filename: 文件名（不含路径，会保存在 AerieOffice 目录下）
        content: 文档内容
        format: 文档格式（markdown / txt）

    Returns:
        保存结果（路径、大小、格式）
    """
    try:
        ext_map = {
            "markdown": ".md",
            "md": ".md",
            "txt": ".txt",
            "text": ".txt",
        }
        ext = ext_map.get(format.lower(), ".md")

        # 确保文件名安全
        safe_name = "".join(c for c in filename if c.isalnum() or c in "._-  ")
        if not safe_name:
            safe_name = f"document_{int(datetime.now().timestamp())}"
        if not safe_name.endswith(ext):
            safe_name += ext

        office_dir = get_office_dir()
        file_path = office_dir / safe_name
        file_path.parent.mkdir(parents=True, exist_ok=True)

        file_path.write_text(content, encoding="utf-8")
        size = file_path.stat().st_size

        return {
            "success": True,
            "path": str(file_path),
            "filename": safe_name,
            "format": format,
            "size_bytes": size,
            "message": f"文档已保存到 {file_path}",
        }
    except Exception as e:
        logger.exception("document_create error")
        return {"success": False, "error": str(e)}


def tool_document_read(filepath: str) -> dict:
    """读取本地文档内容。

    Args:
        filepath: 文件路径（绝对路径或相对 AerieOffice 的路径）

    Returns:
        文档内容和元信息
    """
    try:
        office_dir = get_office_dir()
        path = Path(filepath)
        if not path.is_absolute():
            path = office_dir / filepath

        if not path.exists():
            return {"success": False, "error": f"文件不存在: {path}"}

        if not path.is_file():
            return {"success": False, "error": f"不是文件: {path}"}

        # 安全检查：只允许读取办公目录和常见文档目录
        allowed_parents = [
            office_dir.resolve(),
            Path(os.path.expanduser("~/Desktop")).resolve(),
            Path(os.path.expanduser("~/Documents")).resolve(),
            Path(os.path.expanduser("~/Downloads")).resolve(),
        ]
        resolved = path.resolve()
        allowed = any(str(resolved).startswith(str(p)) for p in allowed_parents)
        if not allowed:
            return {"success": False, "error": "出于安全考虑，仅允许读取桌面/文档/下载/AerieOffice 目录下的文件"}

        # 大小限制：10MB
        size = path.stat().st_size
        if size > 10 * 1024 * 1024:
            return {"success": False, "error": f"文件过大（{size/1024/1024:.1f}MB），最大支持 10MB"}

        content = path.read_text(encoding="utf-8", errors="ignore")

        # 截断超长内容
        truncated = False
        if len(content) > 8000:
            content = content[:8000] + "\n\n... [内容已截断，完整内容请查看原文件]"
            truncated = True

        return {
            "success": True,
            "path": str(resolved),
            "filename": path.name,
            "size_bytes": size,
            "content": content,
            "truncated": truncated,
            "line_count": content.count("\n") + 1,
        }
    except Exception as e:
        logger.exception("document_read error")
        return {"success": False, "error": str(e)}


def tool_spreadsheet_analyze(filepath: str, max_rows: int = 100) -> dict:
    """分析 CSV / TSV 表格文件，输出统计摘要。

    Args:
        filepath: CSV/TSV 文件路径
        max_rows: 最多分析的行数

    Returns:
        列名、行数、每列统计（数值列的均值/最大最小值）
    """
    try:
        office_dir = get_office_dir()
        path = Path(filepath)
        if not path.is_absolute():
            path = office_dir / filepath

        if not path.exists():
            return {"success": False, "error": f"文件不存在: {path}"}

        # 检测分隔符
        ext = path.suffix.lower()
        delimiter = ","
        if ext == ".tsv":
            delimiter = "\t"
        elif ext == ".csv":
            # 嗅探一下
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                first_line = f.readline()
                if "\t" in first_line and "," not in first_line:
                    delimiter = "\t"

        rows: list[dict] = []
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.DictReader(f, delimiter=delimiter)
            for i, row in enumerate(reader):
                if i >= max_rows:
                    break
                rows.append(row)

        if not rows:
            return {"success": False, "error": "表格为空或无法解析"}

        columns = list(rows[0].keys())
        col_stats: dict[str, dict] = {}

        for col in columns:
            values = [row.get(col, "") for row in rows]
            non_empty = [v for v in values if v.strip()]

            # 尝试识别数值列
            numeric_values = []
            for v in non_empty:
                try:
                    numeric_values.append(float(v))
                except (ValueError, TypeError):
                    pass

            if len(numeric_values) >= len(non_empty) * 0.7 and len(numeric_values) > 0:
                col_stats[col] = {
                    "type": "numeric",
                    "count": len(non_empty),
                    "mean": round(sum(numeric_values) / len(numeric_values), 4),
                    "min": min(numeric_values),
                    "max": max(numeric_values),
                }
            else:
                unique_vals = list(set(non_empty))
                col_stats[col] = {
                    "type": "text",
                    "count": len(non_empty),
                    "unique_count": len(unique_vals),
                    "sample_values": unique_vals[:5],
                }

        return {
            "success": True,
            "path": str(path),
            "total_rows": len(rows),
            "columns": columns,
            "column_count": len(columns),
            "column_stats": col_stats,
            "preview": rows[:5],
        }
    except Exception as e:
        logger.exception("spreadsheet_analyze error")
        return {"success": False, "error": str(e)}


def tool_file_search(
    keyword: str,
    directory: str = "",
    file_type: str = "",
    max_results: int = 20,
) -> dict:
    """在本地搜索文件（按文件名关键词）。

    Args:
        keyword: 文件名关键词
        directory: 搜索目录（默认：桌面 + 文档 + 下载 + AerieOffice）
        file_type: 限定文件类型（doc/excel/ppt/pdf/image/code/all）
        max_results: 最多返回结果数

    Returns:
        匹配的文件列表
    """
    try:
        # 搜索范围
        search_dirs = []
        if directory:
            p = Path(directory)
            if p.exists():
                search_dirs.append(p.resolve())
        else:
            for d in ["~/Desktop", "~/Documents", "~/Downloads"]:
                p = Path(os.path.expanduser(d))
                if p.exists():
                    search_dirs.append(p.resolve())
            search_dirs.append(get_office_dir().resolve())

        # 文件类型过滤
        ext_map = {
            "doc": [".md", ".txt", ".doc", ".docx", ".pdf"],
            "excel": [".csv", ".tsv", ".xls", ".xlsx"],
            "ppt": [".ppt", ".pptx"],
            "pdf": [".pdf"],
            "image": [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"],
            "code": [".py", ".js", ".ts", ".html", ".css", ".java", ".go", ".rs", ".cpp"],
            "all": [],
        }
        allowed_exts = ext_map.get(file_type.lower(), [])

        results = []
        keyword_lower = keyword.lower()

        for search_dir in search_dirs:
            if not search_dir.exists():
                continue
            for root, dirs, files in os.walk(search_dir):
                # 跳过隐藏目录
                dirs[:] = [d for d in dirs if not d.startswith(".")]

                for fname in files:
                    if keyword_lower not in fname.lower():
                        continue

                    fpath = Path(root) / fname

                    # 扩展名过滤
                    if allowed_exts and fpath.suffix.lower() not in allowed_exts:
                        continue

                    try:
                        stat = fpath.stat()
                        results.append({
                            "name": fname,
                            "path": str(fpath),
                            "size_bytes": stat.st_size,
                            "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
                        })
                    except Exception:
                        continue

                    if len(results) >= max_results:
                        break
                if len(results) >= max_results:
                    break
            if len(results) >= max_results:
                break

        return {
            "success": True,
            "keyword": keyword,
            "count": len(results),
            "files": results,
            "search_dirs": [str(d) for d in search_dirs],
        }
    except Exception as e:
        logger.exception("file_search error")
        return {"success": False, "error": str(e)}


def tool_text_summary(text: str, max_length: int = 300) -> dict:
    """对文本进行简要摘要（抽取式）。

    Args:
        text: 要摘要的文本
        max_length: 摘要最大字数

    Returns:
        摘要结果
    """
    try:
        if not text or not text.strip():
            return {"success": False, "error": "文本为空"}

        # 简单的抽取式摘要：取前几句话 + 关键词
        sentences = [s.strip() for s in text.replace("。", "。\n").replace("！", "！\n").replace("？", "？\n").split("\n") if s.strip()]

        summary_parts = []
        current_len = 0
        for sent in sentences:
            if current_len + len(sent) > max_length and summary_parts:
                break
            summary_parts.append(sent)
            current_len += len(sent)

        summary = "".join(summary_parts)
        if len(text) > len(summary) and not summary.endswith("。"):
            summary += "..."

        # 提取高频词（简单版）
        words = []
        for char in text:
            if "\u4e00" <= char <= "\u9fff":
                words.append(char)
        word_freq: dict[str, int] = {}
        for i in range(len(words) - 1):
            bigram = words[i] + words[i + 1]
            word_freq[bigram] = word_freq.get(bigram, 0) + 1
        top_keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:5]

        return {
            "success": True,
            "summary": summary,
            "original_length": len(text),
            "summary_length": len(summary),
            "compression_ratio": round(len(summary) / len(text), 3) if text else 0,
            "top_keywords": [kw for kw, _ in top_keywords],
        }
    except Exception as e:
        logger.exception("text_summary error")
        return {"success": False, "error": str(e)}


def tool_calendar_list(
    start_date: str = "",
    end_date: str = "",
    max_results: int = 20,
) -> dict:
    """获取日历事件列表。

    Args:
        start_date: 开始日期 YYYY-MM-DD（默认今天）
        end_date: 结束日期 YYYY-MM-DD（默认 7 天后）
        max_results: 最多返回数量

    Returns:
        事件列表
    """
    try:
        from core.calendar_db import get_calendar_db
        db = get_calendar_db()

        if not start_date:
            start_date = datetime.now().strftime("%Y-%m-%d")
        if not end_date:
            from datetime import timedelta
            end_date = (datetime.now() + timedelta(days=7)).strftime("%Y-%m-%d")

        events = db.list_events(start_date=start_date, end_date=end_date, limit=max_results)
        return {
            "success": True,
            "start_date": start_date,
            "end_date": end_date,
            "count": len(events),
            "events": events,
        }
    except Exception as e:
        logger.exception("calendar_list error")
        return {"success": False, "error": str(e)}


def tool_calendar_create(
    title: str,
    date: str,
    time: str = "",
    description: str = "",
    category: str = "work",
) -> dict:
    """创建日历事件。

    Args:
        title: 事件标题
        date: 日期 YYYY-MM-DD
        time: 时间 HH:MM（可选）
        description: 描述（可选）
        category: 分类（work/personal/reminder 等）

    Returns:
        创建结果
    """
    try:
        from core.calendar_db import get_calendar_db
        db = get_calendar_db()

        event_data = {
            "title": title,
            "date": date,
            "time": time or None,
            "description": description or None,
            "category": category,
        }

        event_id = db.create_event(event_data)
        event = db.get_event(event_id)

        return {
            "success": True,
            "event_id": event_id,
            "event": event,
            "message": f"已创建事件：{title}",
        }
    except Exception as e:
        logger.exception("calendar_create error")
        return {"success": False, "error": str(e)}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 【文件管理类】新增工具
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


def tool_directory_list(
    directory: str = "",
    show_hidden: bool = False,
    max_items: int = 100,
) -> dict:
    """列出目录内容。

    Args:
        directory: 目录路径（默认 AerieOffice）
        show_hidden: 是否显示隐藏文件
        max_items: 最大返回数量

    Returns:
        目录文件列表
    """
    try:
        office_dir = get_office_dir()
        if directory:
            target = Path(directory)
            if not target.is_absolute():
                target = office_dir / directory
        else:
            target = office_dir

        if not target.exists():
            return {"success": False, "error": f"目录不存在: {target}"}
        if not target.is_dir():
            return {"success": False, "error": f"不是目录: {target}"}

        items = []
        for item in target.iterdir():
            if not show_hidden and item.name.startswith("."):
                continue
            try:
                stat = item.stat()
                items.append({
                    "name": item.name,
                    "path": str(item),
                    "is_dir": item.is_dir(),
                    "size_bytes": stat.st_size if item.is_file() else 0,
                    "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
                })
            except Exception:
                continue
            if len(items) >= max_items:
                break

        items.sort(key=lambda x: (not x["is_dir"], x["name"].lower()))

        return {
            "success": True,
            "path": str(target),
            "total_count": len(items),
            "items": items,
        }
    except Exception as e:
        logger.exception("directory_list error")
        return {"success": False, "error": str(e)}


def tool_file_copy(source: str, destination: str) -> dict:
    """复制文件或目录。

    Args:
        source: 源路径
        destination: 目标路径

    Returns:
        复制结果
    """
    try:
        src = Path(source)
        dst = Path(destination)

        if not src.exists():
            return {"success": False, "error": f"源不存在: {source}"}

        if src.is_file():
            shutil.copy2(src, dst)
            size = dst.stat().st_size if dst.exists() else 0
            return {
                "success": True,
                "source": str(src),
                "destination": str(dst),
                "size_bytes": size,
                "message": f"文件已复制到 {dst}",
            }
        elif src.is_dir():
            if dst.exists():
                dst = dst / src.name
            shutil.copytree(src, dst)
            return {
                "success": True,
                "source": str(src),
                "destination": str(dst),
                "message": f"目录已复制到 {dst}",
            }
        else:
            return {"success": False, "error": "不支持的文件类型"}
    except Exception as e:
        logger.exception("file_copy error")
        return {"success": False, "error": str(e)}


def tool_file_move(source: str, destination: str) -> dict:
    """移动文件或目录。

    Args:
        source: 源路径
        destination: 目标路径

    Returns:
        移动结果
    """
    try:
        src = Path(source)
        dst = Path(destination)

        if not src.exists():
            return {"success": False, "error": f"源不存在: {source}"}

        shutil.move(str(src), str(dst))
        return {
            "success": True,
            "source": str(src),
            "destination": str(dst),
            "message": f"已移动到 {dst}",
        }
    except Exception as e:
        logger.exception("file_move error")
        return {"success": False, "error": str(e)}


def tool_file_rename(filepath: str, new_name: str) -> dict:
    """重命名文件。

    Args:
        filepath: 文件路径
        new_name: 新文件名

    Returns:
        重命名结果
    """
    try:
        path = Path(filepath)
        if not path.exists():
            return {"success": False, "error": f"文件不存在: {filepath}"}

        new_path = path.parent / new_name
        path.rename(new_path)
        return {
            "success": True,
            "old_path": str(path),
            "new_path": str(new_path),
            "message": f"已重命名为 {new_name}",
        }
    except Exception as e:
        logger.exception("file_rename error")
        return {"success": False, "error": str(e)}


def tool_directory_create(directory: str) -> dict:
    """创建新目录。

    Args:
        directory: 目录路径

    Returns:
        创建结果
    """
    try:
        office_dir = get_office_dir()
        path = Path(directory)
        if not path.is_absolute():
            path = office_dir / directory

        path.mkdir(parents=True, exist_ok=True)
        return {
            "success": True,
            "path": str(path),
            "message": f"目录已创建: {path}",
        }
    except Exception as e:
        logger.exception("directory_create error")
        return {"success": False, "error": str(e)}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 【文档处理类】新增工具
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


def tool_document_convert(
    filepath: str,
    output_format: str = "markdown",
) -> dict:
    """文档格式转换（txt ↔ md 等简单转换）。

    Args:
        filepath: 输入文件路径
        output_format: 输出格式（markdown / txt / json）

    Returns:
        转换结果
    """
    try:
        path = Path(filepath)
        if not path.exists():
            return {"success": False, "error": f"文件不存在: {filepath}"}

        content = path.read_text(encoding="utf-8", errors="ignore")

        ext_map = {"markdown": ".md", "md": ".md", "txt": ".txt", "json": ".json"}
        out_ext = ext_map.get(output_format.lower(), ".md")
        out_path = path.with_suffix(out_ext)

        if output_format.lower() == "json":
            import json as _json
            out_content = _json.dumps({"content": content, "source": str(path)}, ensure_ascii=False, indent=2)
        else:
            out_content = content

        out_path.write_text(out_content, encoding="utf-8")
        return {
            "success": True,
            "input": str(path),
            "output": str(out_path),
            "output_format": output_format,
            "size_bytes": len(out_content.encode("utf-8")),
        }
    except Exception as e:
        logger.exception("document_convert error")
        return {"success": False, "error": str(e)}


def tool_word_generate(
    filename: str,
    content: str,
    title: str = "",
) -> dict:
    """生成 Word 文档（.docx）。如果未安装 python-docx 则回退为 Markdown。

    Args:
        filename: 文件名（不含路径和后缀）
        content: 文档内容
        title: 标题

    Returns:
        生成结果
    """
    try:
        safe_name = "".join(c for c in filename if c.isalnum() or c in "._-  ") or f"doc_{int(datetime.now().timestamp())}"
        office_dir = get_office_dir()
        file_path = office_dir / f"{safe_name}.docx"

        try:
            from docx import Document
            doc = Document()
            if title:
                doc.add_heading(title, level=0)
            for para in content.split("\n\n"):
                para = para.strip()
                if para:
                    if para.startswith("# "):
                        doc.add_heading(para[2:], level=1)
                    elif para.startswith("## "):
                        doc.add_heading(para[3:], level=2)
                    elif para.startswith("### "):
                        doc.add_heading(para[4:], level=3)
                    elif para.startswith("- ") or para.startswith("* "):
                        doc.add_paragraph(para[2:], style="List Bullet")
                    else:
                        doc.add_paragraph(para)
            doc.save(str(file_path))
            fmt = "docx"
        except ImportError:
            file_path = office_dir / f"{safe_name}.md"
            md_content = (f"# {title}\n\n" if title else "") + content
            file_path.write_text(md_content, encoding="utf-8")
            fmt = "markdown (python-docx 未安装)"

        size = file_path.stat().st_size
        return {
            "success": True,
            "path": str(file_path),
            "filename": file_path.name,
            "format": fmt,
            "size_bytes": size,
            "message": f"文档已保存到 {file_path}",
        }
    except Exception as e:
        logger.exception("word_generate error")
        return {"success": False, "error": str(e)}


def tool_csv_generate(
    filename: str,
    rows: list[dict],
    columns: list[str] | None = None,
) -> dict:
    """生成 CSV 表格文件。

    Args:
        filename: 文件名（不含路径）
        rows: 数据行（字典列表）
        columns: 列名列表（可选，默认从第一行提取）

    Returns:
        生成结果
    """
    try:
        if not rows:
            return {"success": False, "error": "数据为空"}

        safe_name = "".join(c for c in filename if c.isalnum() or c in "._-  ") or f"data_{int(datetime.now().timestamp())}"
        if not safe_name.endswith(".csv"):
            safe_name += ".csv"
        office_dir = get_office_dir()
        file_path = office_dir / safe_name

        cols = columns or list(rows[0].keys())

        with open(file_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=cols, extrasaction="ignore")
            writer.writeheader()
            for row in rows:
                writer.writerow(row)

        size = file_path.stat().st_size
        return {
            "success": True,
            "path": str(file_path),
            "filename": safe_name,
            "row_count": len(rows),
            "column_count": len(cols),
            "columns": cols,
            "size_bytes": size,
        }
    except Exception as e:
        logger.exception("csv_generate error")
        return {"success": False, "error": str(e)}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 【系统操作类】新增工具
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


def tool_system_info() -> dict:
    """获取系统基本信息。

    Returns:
        系统信息
    """
    try:
        import psutil
        office_dir = get_office_dir()
        mem = psutil.virtual_memory()
        disk = psutil.disk_usage(str(office_dir))
        cpu_percent = psutil.cpu_percent(interval=0.1)
        info = {
            "os": platform.system(),
            "os_version": platform.version(),
            "os_release": platform.release(),
            "architecture": platform.machine(),
            "processor": platform.processor() or "未知",
            "hostname": platform.node(),
            "python_version": sys.version,
            "cpu_cores": os.cpu_count(),
            "cpu_usage_percent": cpu_percent,
            "memory_total_gb": round(mem.total / 1024 / 1024 / 1024, 2),
            "memory_used_gb": round(mem.used / 1024 / 1024 / 1024, 2),
            "memory_available_gb": round(mem.available / 1024 / 1024 / 1024, 2),
            "memory_usage_percent": mem.percent,
            "disk_total_gb": round(disk.total / 1024 / 1024 / 1024, 2),
            "disk_used_gb": round(disk.used / 1024 / 1024 / 1024, 2),
            "disk_free_gb": round(disk.free / 1024 / 1024 / 1024, 2),
            "disk_usage_percent": disk.percent,
            "office_dir": str(office_dir),
            "current_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        return {"success": True, "info": info}
    except ImportError:
        office_dir = get_office_dir()
        info = {
            "os": platform.system(),
            "os_version": platform.version(),
            "architecture": platform.machine(),
            "processor": platform.processor() or "未知",
            "hostname": platform.node(),
            "python_version": sys.version,
            "cpu_cores": os.cpu_count(),
            "office_dir": str(office_dir),
            "current_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "note": "psutil 未安装，内存/磁盘/CPU 信息不可用",
        }
        return {"success": True, "info": info}
    except Exception as e:
        logger.exception("system_info error")
        return {"success": False, "error": str(e)}


def tool_process_list(
    keyword: str = "",
    max_results: int = 30,
) -> dict:
    """列出正在运行的进程。

    Args:
        keyword: 按名称关键词过滤
        max_results: 最多返回数量

    Returns:
        进程列表
    """
    try:
        import psutil
        processes = []
        for proc in psutil.process_iter(["pid", "name", "memory_percent", "cpu_percent"]):
            try:
                info = proc.info
                name = info.get("name", "")
                if keyword and keyword.lower() not in name.lower():
                    continue
                processes.append({
                    "pid": info.get("pid"),
                    "name": name,
                    "memory_percent": round(info.get("memory_percent", 0), 2),
                    "cpu_percent": round(info.get("cpu_percent", 0), 2),
                })
                if len(processes) >= max_results:
                    break
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                continue

        processes.sort(key=lambda x: x.get("memory_percent", 0), reverse=True)
        return {
            "success": True,
            "count": len(processes),
            "processes": processes,
        }
    except ImportError:
        return {"success": False, "error": "需要 psutil 库才能查看进程列表"}
    except Exception as e:
        logger.exception("process_list error")
        return {"success": False, "error": str(e)}


def tool_app_open(app_name: str) -> dict:
    """打开应用程序。

    Args:
        app_name: 应用名称（notepad/calc/chrome/edge/word/excel 等快捷名，或完整路径）

    Returns:
        打开结果
    """
    try:
        import subprocess
        app_map = {
            "notepad": "notepad.exe",
            "calc": "calc.exe",
            "calculator": "calc.exe",
            "chrome": "chrome.exe",
            "edge": "msedge.exe",
            "explorer": "explorer.exe",
            "cmd": "cmd.exe",
            "powershell": "powershell.exe",
            "word": "winword.exe",
            "excel": "excel.exe",
            "powerpoint": "powerpnt.exe",
            "ppt": "powerpnt.exe",
            "outlook": "outlook.exe",
        }
        exe = app_map.get(app_name.lower(), app_name)

        if os.path.isabs(exe):
            subprocess.Popen(exe)
        else:
            subprocess.Popen(["cmd", "/c", "start", "", exe], shell=False)

        return {
            "success": True,
            "app": app_name,
            "exe": exe,
            "message": f"正在打开 {app_name}",
        }
    except Exception as e:
        logger.exception("app_open error")
        return {"success": False, "error": str(e)}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 【数据分析类】新增工具
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


def tool_data_stats(
    data: list[dict],
    numeric_columns: list[str] | None = None,
) -> dict:
    """对数据集进行统计分析。

    Args:
        data: 数据集（字典列表）
        numeric_columns: 指定数值列（可选，自动检测）

    Returns:
        统计结果
    """
    try:
        if not data:
            return {"success": False, "error": "数据为空"}

        columns = list(data[0].keys())
        stats: dict[str, dict] = {}

        for col in columns:
            values = [row.get(col) for row in data if row.get(col) is not None and row.get(col) != ""]
            numeric_vals = []
            for v in values:
                try:
                    numeric_vals.append(float(v))
                except (ValueError, TypeError):
                    pass

            if numeric_vals and len(numeric_vals) >= len(values) * 0.5:
                stats[col] = {
                    "type": "numeric",
                    "count": len(numeric_vals),
                    "sum": round(sum(numeric_vals), 4),
                    "mean": round(sum(numeric_vals) / len(numeric_vals), 4),
                    "min": min(numeric_vals),
                    "max": max(numeric_vals),
                    "median": round(sorted(numeric_vals)[len(numeric_vals) // 2], 4),
                }
            else:
                unique = list(set(values))
                stats[col] = {
                    "type": "text",
                    "count": len(values),
                    "unique_count": len(unique),
                    "sample_values": unique[:5],
                }

        return {
            "success": True,
            "row_count": len(data),
            "column_count": len(columns),
            "columns": columns,
            "stats": stats,
        }
    except Exception as e:
        logger.exception("data_stats error")
        return {"success": False, "error": str(e)}


def tool_data_filter(
    data: list[dict],
    column: str,
    operator: str = "equals",
    value: Any = None,
) -> dict:
    """过滤数据集。

    Args:
        data: 数据集
        column: 过滤列名
        operator: 操作符（equals/not_equals/greater/less/contains/between）
        value: 过滤值

    Returns:
        过滤后的数据
    """
    try:
        if not data:
            return {"success": False, "error": "数据为空"}

        result = []
        for row in data:
            cell = row.get(column)
            try:
                if operator == "equals":
                    if str(cell) == str(value):
                        result.append(row)
                elif operator == "not_equals":
                    if str(cell) != str(value):
                        result.append(row)
                elif operator == "greater":
                    if cell is not None and float(cell) > float(value):
                        result.append(row)
                elif operator == "less":
                    if cell is not None and float(cell) < float(value):
                        result.append(row)
                elif operator == "contains":
                    if value and str(value).lower() in str(cell).lower():
                        result.append(row)
                elif operator == "between" and isinstance(value, (list, tuple)) and len(value) == 2:
                    if cell is not None and float(value[0]) <= float(cell) <= float(value[1]):
                        result.append(row)
            except (ValueError, TypeError):
                continue

        return {
            "success": True,
            "original_count": len(data),
            "filtered_count": len(result),
            "data": result,
        }
    except Exception as e:
        logger.exception("data_filter error")
        return {"success": False, "error": str(e)}


def tool_data_sort(
    data: list[dict],
    column: str,
    ascending: bool = True,
) -> dict:
    """对数据集进行排序。

    Args:
        data: 数据集
        column: 排序列名
        ascending: 是否升序

    Returns:
        排序后的数据
    """
    try:
        if not data:
            return {"success": False, "error": "数据为空"}

        def sort_key(row):
            v = row.get(column)
            try:
                return (0, float(v))
            except (ValueError, TypeError):
                return (1, str(v) if v else "")

        sorted_data = sorted(data, key=sort_key, reverse=not ascending)
        return {
            "success": True,
            "count": len(sorted_data),
            "column": column,
            "order": "ascending" if ascending else "descending",
            "data": sorted_data,
        }
    except Exception as e:
        logger.exception("data_sort error")
        return {"success": False, "error": str(e)}


def tool_chart_generate(
    data: list[dict],
    x_column: str,
    y_column: str,
    chart_type: str = "bar",
    title: str = "",
    width: int = 600,
    height: int = 400,
) -> dict:
    """生成 SVG 图表（柱状图 / 折线图 / 饼图）。

    Args:
        data: 数据集
        x_column: X 轴列名
        y_column: Y 轴列名
        chart_type: 图表类型（bar/line/pie）
        title: 图表标题
        width: 宽度
        height: 高度

    Returns:
        SVG 图表内容
    """
    try:
        if not data:
            return {"success": False, "error": "数据为空"}

        labels = [str(row.get(x_column, "")) for row in data]
        values = []
        for row in data:
            try:
                values.append(float(row.get(y_column, 0)))
            except (ValueError, TypeError):
                values.append(0)

        max_val = max(values) if values else 1
        min_val = 0
        padding = {"top": 50, "right": 30, "bottom": 60, "left": 60}
        chart_w = width - padding["left"] - padding["right"]
        chart_h = height - padding["top"] - padding["bottom"]

        svg_parts = [
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
            '<rect width="100%" height="100%" fill="#fff"/>',
        ]

        if title:
            svg_parts.append(f'<text x="{width/2}" y="28" text-anchor="middle" font-size="16" font-weight="bold" fill="#333">{title}</text>')

        if chart_type == "bar":
            bar_w = chart_w / len(data) * 0.7
            gap = chart_w / len(data) * 0.3
            for i, (label, val) in enumerate(zip(labels, values)):
                x = padding["left"] + i * (bar_w + gap) + gap / 2
                bar_h = (val / max_val) * chart_h if max_val > 0 else 0
                y = padding["top"] + chart_h - bar_h
                color = f"hsl({(i * 360 / len(data)) % 360}, 70%, 60%)"
                svg_parts.append(f'<rect x="{x:.1f}" y="{y:.1f}" width="{bar_w:.1f}" height="{bar_h:.1f}" fill="{color}" rx="3"/>')
                svg_parts.append(f'<text x="{x + bar_w/2:.1f}" y="{y - 5}" text-anchor="middle" font-size="11" fill="#666">{val}</text>')
                svg_parts.append(f'<text x="{x + bar_w/2:.1f}" y="{padding["top"] + chart_h + 18}" text-anchor="middle" font-size="11" fill="#666" transform="rotate(-25 {x + bar_w/2:.1f},{padding["top"] + chart_h + 18})">{label[:10]}</text>')

        elif chart_type == "line":
            points = []
            for i, (label, val) in enumerate(zip(labels, values)):
                x = padding["left"] + (i / max(len(data) - 1, 1)) * chart_w
                y = padding["top"] + chart_h - (val / max_val) * chart_h if max_val > 0 else padding["top"] + chart_h
                points.append(f"{x:.1f},{y:.1f}")
            path = "M" + " L".join(points)
            svg_parts.append(f'<polyline points="{" ".join(points)}" fill="none" stroke="#4f8cff" stroke-width="2"/>')
            for i, (label, val) in enumerate(zip(labels, values)):
                x = padding["left"] + (i / max(len(data) - 1, 1)) * chart_w
                y = padding["top"] + chart_h - (val / max_val) * chart_h if max_val > 0 else padding["top"] + chart_h
                svg_parts.append(f'<circle cx="{x:.1f}" cy="{y:.1f}" r="4" fill="#4f8cff"/>')
                svg_parts.append(f'<text x="{x:.1f}" y="{y - 8}" text-anchor="middle" font-size="10" fill="#666">{val}</text>')

        elif chart_type == "pie":
            cx, cy, r = width / 2, height / 2 + 10, min(chart_w, chart_h) / 2 - 10
            total = sum(values) or 1
            start_angle = 0
            for i, (label, val) in enumerate(zip(labels, values)):
                angle = (val / total) * 360
                end_angle = start_angle + angle
                color = f"hsl({(i * 360 / len(data)) % 360}, 70%, 60%)"
                import math
                x1 = cx + r * math.cos(math.radians(start_angle - 90))
                y1 = cy + r * math.sin(math.radians(start_angle - 90))
                x2 = cx + r * math.cos(math.radians(end_angle - 90))
                y2 = cy + r * math.sin(math.radians(end_angle - 90))
                large_arc = 1 if angle > 180 else 0
                svg_parts.append(f'<path d="M{cx},{cy} L{x1:.1f},{y1:.1f} A{r},{r} 0 {large_arc},1 {x2:.1f},{y2:.1f} Z" fill="{color}"/>')
                mid_angle = start_angle + angle / 2
                lx = cx + (r * 0.6) * math.cos(math.radians(mid_angle - 90))
                ly = cy + (r * 0.6) * math.sin(math.radians(mid_angle - 90))
                if angle > 15:
                    svg_parts.append(f'<text x="{lx:.1f}" y="{ly:.1f}" text-anchor="middle" font-size="10" fill="#fff">{label[:6]}</text>')
                start_angle = end_angle

        # 坐标轴
        if chart_type in ("bar", "line"):
            svg_parts.append(f'<line x1="{padding["left"]}" y1="{padding["top"]}" x2="{padding["left"]}" y2="{padding["top"] + chart_h}" stroke="#ccc"/>')
            svg_parts.append(f'<line x1="{padding["left"]}" y1="{padding["top"] + chart_h}" x2="{padding["left"] + chart_w}" y2="{padding["top"] + chart_h}" stroke="#ccc"/>')

        svg_parts.append("</svg>")
        svg_content = "\n".join(svg_parts)

        # 保存到文件
        safe_title = "".join(c for c in title if c.isalnum() or c in "._-  ") or "chart"
        office_dir = get_office_dir()
        chart_path = office_dir / f"{safe_title}.svg"
        chart_path.write_text(svg_content, encoding="utf-8")

        return {
            "success": True,
            "chart_type": chart_type,
            "title": title or "chart",
            "data_points": len(data),
            "svg_path": str(chart_path),
            "svg_content": svg_content[:2000] + "..." if len(svg_content) > 2000 else svg_content,
            "message": f"图表已保存到 {chart_path}",
        }
    except Exception as e:
        logger.exception("chart_generate error")
        return {"success": False, "error": str(e)}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 【网络工具类】新增工具
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


def tool_web_fetch(url: str, max_length: int = 5000) -> dict:
    """抓取网页内容（纯文本）。

    Args:
        url: 网页 URL
        max_length: 最大返回字符数

    Returns:
        网页内容
    """
    try:
        import urllib.request
        from urllib.parse import urlparse

        parsed = urlparse(url)
        if not parsed.scheme:
            url = "https://" + url

        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 AerieOffice/13.9"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            raw = resp.read()

        # 尝试 UTF-8，不行用 GBK
        try:
            content = raw.decode("utf-8", errors="ignore")
        except Exception:
            content = raw.decode("gbk", errors="ignore")

        # 简单去除 HTML 标签
        import re
        text = re.sub(r"<script.*?</script>", "", content, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).strip()

        truncated = len(text) > max_length
        if truncated:
            text = text[:max_length] + "..."

        return {
            "success": True,
            "url": url,
            "content": text,
            "original_length": len(text),
            "truncated": truncated,
        }
    except Exception as e:
        logger.exception("web_fetch error")
        return {"success": False, "error": str(e)}


def tool_weather_query(city: str = "") -> dict:
    """查询天气（调用 wttr.in 免费接口）。

    Args:
        city: 城市名（中文/英文）

    Returns:
        天气信息
    """
    try:
        import urllib.request
        import urllib.parse

        if not city:
            city = "auto"

        encoded_city = urllib.parse.quote(city)
        url = f"https://wttr.in/{encoded_city}?format=j1"
        req = urllib.request.Request(url, headers={"User-Agent": "curl/8.0"})

        with urllib.request.urlopen(req, timeout=8) as resp:
            data = json.loads(resp.read().decode("utf-8"))

        current = data.get("current_condition", [{}])[0]
        weather_desc = current.get("lang_zh", [{}])[0].get("value", current.get("weatherDesc", [{}])[0].get("value", ""))
        temp_c = current.get("temp_C", "")
        feels_like = current.get("FeelsLikeC", "")
        humidity = current.get("humidity", "")
        wind_speed = current.get("windspeedKmph", "")
        wind_dir = current.get("winddir16Point", "")

        result = {
            "city": city,
            "description": weather_desc,
            "temperature_c": temp_c,
            "feels_like_c": feels_like,
            "humidity_percent": humidity,
            "wind_speed_kmh": wind_speed,
            "wind_direction": wind_dir,
            "update_time": datetime.now().strftime("%Y-%m-%d %H:%M"),
        }
        return {"success": True, "weather": result}
    except Exception as e:
        logger.exception("weather_query error")
        return {"success": False, "error": str(e)}


def tool_translation(
    text: str,
    target_lang: str = "zh",
    source_lang: str = "auto",
) -> dict:
    """文本翻译（使用免费的 MyMemory API）。

    Args:
        text: 要翻译的文本
        target_lang: 目标语言（zh/en/ja/ko 等）
        source_lang: 源语言（auto 自动检测）

    Returns:
        翻译结果
    """
    try:
        import urllib.request
        import urllib.parse

        lang_map = {
            "zh": "zh-CN", "cn": "zh-CN", "chinese": "zh-CN",
            "en": "en", "english": "en",
            "ja": "ja", "jp": "ja", "japanese": "ja",
            "ko": "ko", "kr": "ko", "korean": "ko",
            "fr": "fr", "de": "de", "es": "es", "ru": "ru",
        }
        tl = lang_map.get(target_lang.lower(), target_lang)
        sl = "autodetect" if source_lang.lower() == "auto" else lang_map.get(source_lang.lower(), source_lang)

        encoded_text = urllib.parse.quote(text[:500])
        url = f"https://api.mymemory.translated.net/get?q={encoded_text}&langpair={sl}|{tl}"

        with urllib.request.urlopen(url, timeout=10) as resp:
            data = json.loads(resp.read().decode("utf-8"))

        translated = data.get("responseData", {}).get("translatedText", "")
        return {
            "success": bool(translated),
            "original_text": text,
            "translated_text": translated,
            "source_lang": sl,
            "target_lang": tl,
        }
    except Exception as e:
        logger.exception("translation error")
        return {"success": False, "error": str(e)}


def tool_code_search(
    query: str,
    language: str = "python",
    max_results: int = 10,
) -> dict:
    """代码搜索（通过 GitHub 公开搜索接口）。

    Args:
        query: 搜索关键词
        language: 编程语言
        max_results: 最多返回数量

    Returns:
        搜索结果
    """
    try:
        import urllib.request
        import urllib.parse

        encoded_q = urllib.parse.quote(f"{query} language:{language}")
        url = f"https://api.github.com/search/code?q={encoded_q}&per_page={max_results}"
        req = urllib.request.Request(url, headers={
            "User-Agent": "AerieOffice/13.9",
            "Accept": "application/vnd.github.v3+json",
        })

        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read().decode("utf-8"))

        items = []
        for item in data.get("items", [])[:max_results]:
            items.append({
                "name": item.get("name"),
                "path": item.get("path"),
                "repo": item.get("repository", {}).get("full_name"),
                "url": item.get("html_url"),
                "language": item.get("language", language),
            })

        return {
            "success": True,
            "query": query,
            "language": language,
            "total_count": data.get("total_count", 0),
            "results": items,
        }
    except Exception as e:
        logger.exception("code_search error")
        return {"success": False, "error": str(e)}


# ── 注册到 ToolRegistry ──────────────────────────

_OFFICE_TOOL_SCHEMAS = {
    "document_create": {
        "type": "function",
        "function": {
            "name": "document_create",
            "description": "创建并保存一个办公文档（Markdown/TXT格式）到本地 AerieOffice 目录。用于写报告、方案、总结、邮件草稿等。",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "文件名（不需要路径和后缀）",
                    },
                    "content": {
                        "type": "string",
                        "description": "文档完整内容",
                    },
                    "format": {
                        "type": "string",
                        "description": "文档格式：markdown 或 txt",
                        "enum": ["markdown", "txt"],
                        "default": "markdown",
                    },
                },
                "required": ["filename", "content"],
            },
        },
    },
    "document_read": {
        "type": "function",
        "function": {
            "name": "document_read",
            "description": "读取本地文档内容（支持 md/txt/csv 等文本文件）。仅允许读取桌面/文档/下载/AerieOffice 目录下的文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "filepath": {
                        "type": "string",
                        "description": "文件路径（绝对路径或相对 AerieOffice 的路径）",
                    },
                },
                "required": ["filepath"],
            },
        },
    },
    "spreadsheet_analyze": {
        "type": "function",
        "function": {
            "name": "spreadsheet_analyze",
            "description": "分析 CSV/TSV 表格文件，输出列名、行数、数值列统计（均值/最大最小）、数据预览。",
            "parameters": {
                "type": "object",
                "properties": {
                    "filepath": {
                        "type": "string",
                        "description": "CSV/TSV 文件路径",
                    },
                    "max_rows": {
                        "type": "integer",
                        "description": "最多分析的行数",
                        "default": 100,
                    },
                },
                "required": ["filepath"],
            },
        },
    },
    "file_search": {
        "type": "function",
        "function": {
            "name": "file_search",
            "description": "在本地搜索文件（按文件名关键词）。默认搜索桌面、文档、下载、AerieOffice 目录。",
            "parameters": {
                "type": "object",
                "properties": {
                    "keyword": {
                        "type": "string",
                        "description": "文件名关键词",
                    },
                    "directory": {
                        "type": "string",
                        "description": "指定搜索目录（留空则搜索默认位置）",
                    },
                    "file_type": {
                        "type": "string",
                        "description": "文件类型过滤",
                        "enum": ["all", "doc", "excel", "ppt", "pdf", "image", "code"],
                        "default": "all",
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "最多返回结果数",
                        "default": 20,
                    },
                },
                "required": ["keyword"],
            },
        },
    },
    "text_summary": {
        "type": "function",
        "function": {
            "name": "text_summary",
            "description": "对长文本进行快速摘要，提取核心内容和关键词。",
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {
                        "type": "string",
                        "description": "要摘要的文本内容",
                    },
                    "max_length": {
                        "type": "integer",
                        "description": "摘要最大字数",
                        "default": 300,
                    },
                },
                "required": ["text"],
            },
        },
    },
    "calendar_list": {
        "type": "function",
        "function": {
            "name": "calendar_list",
            "description": "查看日历事件列表，了解近期日程安排。",
            "parameters": {
                "type": "object",
                "properties": {
                    "start_date": {
                        "type": "string",
                        "description": "开始日期 YYYY-MM-DD（默认今天）",
                    },
                    "end_date": {
                        "type": "string",
                        "description": "结束日期 YYYY-MM-DD（默认 7 天后）",
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "最多返回数量",
                        "default": 20,
                    },
                },
                "required": [],
            },
        },
    },
    "calendar_create": {
        "type": "function",
        "function": {
            "name": "calendar_create",
            "description": "创建日历事件，安排会议、提醒、待办事项等。",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {
                        "type": "string",
                        "description": "事件标题",
                    },
                    "date": {
                        "type": "string",
                        "description": "日期 YYYY-MM-DD",
                    },
                    "time": {
                        "type": "string",
                        "description": "时间 HH:MM（可选）",
                    },
                    "description": {
                        "type": "string",
                        "description": "事件描述（可选）",
                    },
                    "category": {
                        "type": "string",
                        "description": "分类",
                        "enum": ["work", "personal", "reminder", "meeting"],
                        "default": "work",
                    },
                },
                "required": ["title", "date"],
            },
        },
    },
    # ━━━ 文件管理类新增 ━━━
    "directory_list": {
        "type": "function",
        "function": {
            "name": "directory_list",
            "description": "列出指定目录下的文件和文件夹。默认列出 AerieOffice 目录内容。",
            "parameters": {
                "type": "object",
                "properties": {
                    "directory": {
                        "type": "string",
                        "description": "目录路径（留空则为 AerieOffice）",
                    },
                    "show_hidden": {
                        "type": "boolean",
                        "description": "是否显示隐藏文件",
                        "default": False,
                    },
                    "max_items": {
                        "type": "integer",
                        "description": "最大返回数量",
                        "default": 100,
                    },
                },
                "required": [],
            },
        },
    },
    "file_copy": {
        "type": "function",
        "function": {
            "name": "file_copy",
            "description": "复制文件或目录到目标位置。",
            "parameters": {
                "type": "object",
                "properties": {
                    "source": {
                        "type": "string",
                        "description": "源文件或源目录路径",
                    },
                    "destination": {
                        "type": "string",
                        "description": "目标路径",
                    },
                },
                "required": ["source", "destination"],
            },
        },
    },
    "file_move": {
        "type": "function",
        "function": {
            "name": "file_move",
            "description": "移动文件或目录到目标位置。",
            "parameters": {
                "type": "object",
                "properties": {
                    "source": {
                        "type": "string",
                        "description": "源文件或源目录路径",
                    },
                    "destination": {
                        "type": "string",
                        "description": "目标路径",
                    },
                },
                "required": ["source", "destination"],
            },
        },
    },
    "file_rename": {
        "type": "function",
        "function": {
            "name": "file_rename",
            "description": "重命名文件或目录。",
            "parameters": {
                "type": "object",
                "properties": {
                    "filepath": {
                        "type": "string",
                        "description": "原文件路径",
                    },
                    "new_name": {
                        "type": "string",
                        "description": "新名称",
                    },
                },
                "required": ["filepath", "new_name"],
            },
        },
    },
    "directory_create": {
        "type": "function",
        "function": {
            "name": "directory_create",
            "description": "创建新目录（支持递归创建多级目录）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "directory": {
                        "type": "string",
                        "description": "目录路径（相对路径会创建在 AerieOffice 下）",
                    },
                },
                "required": ["directory"],
            },
        },
    },
    # ━━━ 文档处理类新增 ━━━
    "document_convert": {
        "type": "function",
        "function": {
            "name": "document_convert",
            "description": "文档格式转换，支持 markdown / txt / json 之间互转。",
            "parameters": {
                "type": "object",
                "properties": {
                    "filepath": {
                        "type": "string",
                        "description": "输入文件路径",
                    },
                    "output_format": {
                        "type": "string",
                        "description": "输出格式",
                        "enum": ["markdown", "txt", "json"],
                        "default": "markdown",
                    },
                },
                "required": ["filepath"],
            },
        },
    },
    "word_generate": {
        "type": "function",
        "function": {
            "name": "word_generate",
            "description": "生成 Word 文档（.docx 格式）。支持 Markdown 语法解析为标题、列表等。未安装 python-docx 时回退为 Markdown。",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "文件名（不含路径和后缀）",
                    },
                    "content": {
                        "type": "string",
                        "description": "文档内容，支持 Markdown 标题和列表语法",
                    },
                    "title": {
                        "type": "string",
                        "description": "文档标题（可选）",
                    },
                },
                "required": ["filename", "content"],
            },
        },
    },
    "csv_generate": {
        "type": "function",
        "function": {
            "name": "csv_generate",
            "description": "生成 CSV 表格文件，自动处理中文编码（UTF-8 BOM）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "文件名（不含路径）",
                    },
                    "rows": {
                        "type": "array",
                        "description": "数据行列表，每行为字典",
                        "items": {"type": "object"},
                    },
                    "columns": {
                        "type": "array",
                        "description": "列名列表（可选，默认从第一行提取）",
                        "items": {"type": "string"},
                    },
                },
                "required": ["filename", "rows"],
            },
        },
    },
    # ━━━ 系统操作类新增 ━━━
    "system_info": {
        "type": "function",
        "function": {
            "name": "system_info",
            "description": "获取系统基本信息：操作系统、CPU、内存、磁盘、Python 版本等。",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": [],
            },
        },
    },
    "process_list": {
        "type": "function",
        "function": {
            "name": "process_list",
            "description": "列出当前正在运行的进程，可按名称关键词过滤。需要 psutil 库。",
            "parameters": {
                "type": "object",
                "properties": {
                    "keyword": {
                        "type": "string",
                        "description": "进程名关键词过滤（可选）",
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "最多返回数量",
                        "default": 30,
                    },
                },
                "required": [],
            },
        },
    },
    "app_open": {
        "type": "function",
        "function": {
            "name": "app_open",
            "description": "打开应用程序。支持快捷名（notepad/calc/chrome/edge/word/excel 等）或完整路径。",
            "parameters": {
                "type": "object",
                "properties": {
                    "app_name": {
                        "type": "string",
                        "description": "应用名称或路径",
                    },
                },
                "required": ["app_name"],
            },
        },
    },
    # ━━━ 数据分析类新增 ━━━
    "data_stats": {
        "type": "function",
        "function": {
            "name": "data_stats",
            "description": "对数据集进行统计分析：数值列计算均值/求和/最大最小/中位数，文本列统计去重数和样本。",
            "parameters": {
                "type": "object",
                "properties": {
                    "data": {
                        "type": "array",
                        "description": "数据集（字典列表）",
                        "items": {"type": "object"},
                    },
                    "numeric_columns": {
                        "type": "array",
                        "description": "指定数值列（可选，自动检测）",
                        "items": {"type": "string"},
                    },
                },
                "required": ["data"],
            },
        },
    },
    "data_filter": {
        "type": "function",
        "function": {
            "name": "data_filter",
            "description": "按条件过滤数据集，支持等于、不等于、大于、小于、包含、区间等操作。",
            "parameters": {
                "type": "object",
                "properties": {
                    "data": {
                        "type": "array",
                        "description": "数据集（字典列表）",
                        "items": {"type": "object"},
                    },
                    "column": {
                        "type": "string",
                        "description": "过滤列名",
                    },
                    "operator": {
                        "type": "string",
                        "description": "操作符",
                        "enum": ["equals", "not_equals", "greater", "less", "contains", "between"],
                        "default": "equals",
                    },
                    "value": {
                        "description": "过滤值（between 操作时为 [min, max] 数组）",
                    },
                },
                "required": ["data", "column"],
            },
        },
    },
    "data_sort": {
        "type": "function",
        "function": {
            "name": "data_sort",
            "description": "按指定列对数据集排序，支持升序和降序。数值优先排序。",
            "parameters": {
                "type": "object",
                "properties": {
                    "data": {
                        "type": "array",
                        "description": "数据集（字典列表）",
                        "items": {"type": "object"},
                    },
                    "column": {
                        "type": "string",
                        "description": "排序列名",
                    },
                    "ascending": {
                        "type": "boolean",
                        "description": "是否升序",
                        "default": True,
                    },
                },
                "required": ["data", "column"],
            },
        },
    },
    "chart_generate": {
        "type": "function",
        "function": {
            "name": "chart_generate",
            "description": "生成 SVG 格式的图表，支持柱状图、折线图、饼图。图表会保存到 AerieOffice 目录。",
            "parameters": {
                "type": "object",
                "properties": {
                    "data": {
                        "type": "array",
                        "description": "数据集（字典列表）",
                        "items": {"type": "object"},
                    },
                    "x_column": {
                        "type": "string",
                        "description": "X 轴数据列名",
                    },
                    "y_column": {
                        "type": "string",
                        "description": "Y 轴数据列名",
                    },
                    "chart_type": {
                        "type": "string",
                        "description": "图表类型",
                        "enum": ["bar", "line", "pie"],
                        "default": "bar",
                    },
                    "title": {
                        "type": "string",
                        "description": "图表标题（可选）",
                    },
                    "width": {
                        "type": "integer",
                        "description": "图表宽度（像素）",
                        "default": 600,
                    },
                    "height": {
                        "type": "integer",
                        "description": "图表高度（像素）",
                        "default": 400,
                    },
                },
                "required": ["data", "x_column", "y_column"],
            },
        },
    },
    # ━━━ 网络工具类新增 ━━━
    "web_fetch": {
        "type": "function",
        "function": {
            "name": "web_fetch",
            "description": "抓取网页内容并提取纯文本（去除 HTML 标签）。注意：仅用于公开网页，禁止用于需要登录或敏感内容。",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "网页 URL",
                    },
                    "max_length": {
                        "type": "integer",
                        "description": "最大返回字符数",
                        "default": 5000,
                    },
                },
                "required": ["url"],
            },
        },
    },
    "weather_query": {
        "type": "function",
        "function": {
            "name": "weather_query",
            "description": "查询指定城市的当前天气信息（温度、体感、湿度、风速等）。调用免费公开接口。",
            "parameters": {
                "type": "object",
                "properties": {
                    "city": {
                        "type": "string",
                        "description": "城市名（中文或英文，留空自动定位）",
                    },
                },
                "required": [],
            },
        },
    },
    "translation": {
        "type": "function",
        "function": {
            "name": "translation",
            "description": "文本翻译，支持中英日韩法德俄西等多语言互译。调用免费翻译 API。",
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {
                        "type": "string",
                        "description": "要翻译的文本",
                    },
                    "target_lang": {
                        "type": "string",
                        "description": "目标语言代码",
                        "enum": ["zh", "en", "ja", "ko", "fr", "de", "es", "ru"],
                        "default": "zh",
                    },
                    "source_lang": {
                        "type": "string",
                        "description": "源语言（auto 自动检测）",
                        "default": "auto",
                    },
                },
                "required": ["text"],
            },
        },
    },
    "code_search": {
        "type": "function",
        "function": {
            "name": "code_search",
            "description": "在 GitHub 上搜索代码，支持按编程语言过滤。用于查找代码示例和参考实现。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "搜索关键词",
                    },
                    "language": {
                        "type": "string",
                        "description": "编程语言",
                        "enum": ["python", "javascript", "typescript", "java", "go", "rust", "cpp", "csharp", "ruby", "php"],
                        "default": "python",
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "最多返回结果数",
                        "default": 10,
                    },
                },
                "required": ["query"],
            },
        },
    },
}


def register_office_tools(registry) -> int:
    """注册所有办公工具到 ToolRegistry。返回注册的工具数量。"""
    from .tool_registry import ToolRegistry

    tool_funcs = {
        # 文件管理类
        "document_create": tool_document_create,
        "document_read": tool_document_read,
        "file_search": tool_file_search,
        "directory_list": tool_directory_list,
        "file_copy": tool_file_copy,
        "file_move": tool_file_move,
        "file_rename": tool_file_rename,
        "directory_create": tool_directory_create,
        # 文档处理类
        "text_summary": tool_text_summary,
        "document_convert": tool_document_convert,
        "word_generate": tool_word_generate,
        "spreadsheet_analyze": tool_spreadsheet_analyze,
        "csv_generate": tool_csv_generate,
        # 系统操作类
        "calendar_list": tool_calendar_list,
        "calendar_create": tool_calendar_create,
        "system_info": tool_system_info,
        "process_list": tool_process_list,
        "app_open": tool_app_open,
        # 数据分析类
        "data_stats": tool_data_stats,
        "data_filter": tool_data_filter,
        "data_sort": tool_data_sort,
        "chart_generate": tool_chart_generate,
        # 网络工具类
        "web_fetch": tool_web_fetch,
        "weather_query": tool_weather_query,
        "translation": tool_translation,
        "code_search": tool_code_search,
    }

    count = 0
    for name, func in tool_funcs.items():
        schema = _OFFICE_TOOL_SCHEMAS.get(name)
        if schema:
            registry.register(name, func, schema)
            count += 1

    logger.info("registered %d office tools", count)
    return count
